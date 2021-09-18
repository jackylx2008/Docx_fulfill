from numpy.core.numerictypes import _can_coerce_all
import pandas as pd
import os
import shutil
import time
import re
from docx import Document

xlsx_B25B26_columns = ["审批表编号", "提出单位一", "变更原因二",
                       "变更内容三", "变更通知单编号", "变更编号和说明", "日期五", "六版本", "是否完成面单"]
xlsx_B23_columns = ["审批表编号", "提出单位一", "变更原因二", "变更内容三",
                    "变更通知单编号", "变更编号和说明", "日期五",  "是否完成面单", "是否正式盖章蓝图"]


class base_data_from_xls:
    """ # * 用于对输入的Excel文件内的数据读取并进行基础处理
    # * 对于重复的编号进行复制，通过输入的变更单前缀编号生成何种专业，避免出错
    """

    def __init__(self, file_name: str, project_id: str = None):
        self.base_data_no_proc = pd.read_excel(file_name)
        self.base_data = pd.DataFrame()
        if project_id == "B25B26":
            self.base_data_no_proc.columns = xlsx_B25B26_columns
        elif project_id == "B23":
            self.base_data_no_proc.columns = xlsx_B23_columns
        self.base_data_no_proc = self.base_data_no_proc[self.base_data_no_proc.loc[:, "是否完成面单"] == "否"]
        print(self.base_data_no_proc)
        if len(self.base_data_no_proc) == 0:
            print("空列表，单子都出过了")
            exit()

    @staticmethod
    def set_major_cols(x: pd.DataFrame) -> str:
        """根据变更前缀编号确定何种专业
           # ?没想到更简单的方法，所以写了这么一个函数的

        Args:
            x (DataFrame): 取df的变更通知单编号的列

        Returns:
            str: 返回何种专业的str
        """

        if "05-0" in x["变更通知单编号"]:
            return "给排水"
        if "06-0" in x["变更通知单编号"]:
            return "暖通空调"
        if "07-0" in x["变更通知单编号"]:
            return "强电"
        if "R20-" in x["变更通知单编号"]:
            return "热力"
        if "JPS-" in x["变更通知单编号"]:
            return "厨房给排水"
        if "NT-" in x["变更通知单编号"]:
            return "厨房暖通"

    def get_data_proceed(self) -> pd.DataFrame:
        """直接处理输入excel：
            1.检查是否出现编号（包括审批表和变更单）重复，如果重复返回空df
            2.根据审批表内对应的索引生成列，并返回df

        Returns:
            Dataframe:
        """
        # * 检查是否出现重复数据
        flag1 = self.base_data_no_proc["变更通知单编号"].duplicated()
        flag2 = self.base_data_no_proc["审批表编号"].duplicated()
        if (flag1.any() == True or flag2.any() == True):
            repeat_df1 = self.base_data_no_proc["变更通知单编号"].count() > 1
            repeat_df2 = self.base_data_no_proc["审批表编号"].count() > 1
            print(repeat_df1)
            print(repeat_df2)
            print("出现重复编号")
            return pd.DataFrame()
        else:
            self.base_data = self.base_data_no_proc
            self.base_data["何种专业四"] = self.base_data.apply(
                self.set_major_cols, axis=1)
            self.base_data["#####"] = self.base_data["变更通知单编号"]
            self.base_data.drop(columns="变更通知单编号")
            return self.base_data


class replace_doc_lines():
    """
    # * 基于带索引的审批表文件docx文件，自动填充内容并生成带编号的docx
    # * 每个一个生成的docx是一个实例，处理记得save

        Args:
            basic_doc_file (str): 带索引的审批表文件docx文件绝对路径
            doc_dir: str (str): 自动填充生成的docx存放文件夹绝对路径
    """

    def __init__(self, basic_doc_file: str, doc_dir: str, summary: str):
        self.basic_doc_file = basic_doc_file
        self.proceed_doc_dir = doc_dir
        # self.file_num = str(num+1)
        self.file_summary = summary
        self.proceed_doc_file = ""

    def make_new_docx(self, num: int) -> None:
        """基于base文件拷贝生成新的docx文件，文件名带编号内容不做调整

        Args:
            num (int): 审批表编号

        # ? 这类函数是不是应该有个返回值说明是否运行成功?
        Returns:
            无返回值
        """
        # file_num = str(num).rjust(3, '0')  # 在df里面是的类型是numpy.int64，转成str填零
        self.proceed_doc_file = self.proceed_doc_dir + \
            self.basic_doc_file.split(
                '.')[0] + self.file_summary + ".docx"
        shutil.copyfile(self.basic_doc_file, self.proceed_doc_file)
        self.proceed_doc = Document(self.proceed_doc_file)

    def replace_text(self, old_text: str, new_input: str) -> None:
        """替换docx文档中的文字，包括表格里和正文里的

        Args:
            old_text (str): base文件内的索引
            new_input (str): 从df中取出来的内容

        # ? 这类函数是不是应该有个返回值说明是否运行成功?
        Returns:
            无返回值
        """
        if isinstance(new_input, str):
            new_text = new_input
        else:
            # 在df里面是的类型是numpy.int64，转成str填零
            new_text = str(new_input).rjust(3, '0')
        # * 下面这段是从网上抄过来的，因为base文件内需要替换的内容在表格里，所以需要遍历到table->row->cell->paragraph->run（每行）
        for table in self.proceed_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # print(run.text)
                            run.text = run.text.replace(old_text, new_text)
        # * 下面这段也是从网上抄过来的，用替换审批表的编号
        for paragraph in self.proceed_doc.paragraphs:
            for run in paragraph.runs:
                # print(run.text)
                run.text = run.text.replace(old_text, new_text)

    def save_proceed_docx(self):
        self.proceed_doc.save(self.proceed_doc_file)


def copy_file_to_target(src_dir: str, file_key_name: str, des_dir: str) -> True or False:
    """从src_dir拷贝含有特定关键词的文件到des_dir

    Args:
        src_dir (str): 源文件夹
        file_key_name (str): 关键词
        des_dir (str): 目标文件夹

    Returns:
        True or False: 成功在src_dir找到file_key_name并拷贝到des目录下返回True
        ，否则返回False
    """
    if os.path.exists is False:
        os.mkdir(des_dir)
        print("新建 %s " % des_dir)
    # 遍历生成src_folders绝对路径lists
    files_under_src = []
    for dirpath, dirnames, filenames in os.walk(src_dir):
        for filename in filenames:
            files_under_src.append(os.path.join(dirpath, filename))
            print(os.path.join(dirpath, filename))

    for file in files_under_src:
        if str.find(file, file_key_name) >= 0:
            shutil.copy2(file, des_dir)
            print(file + " --> " + des_dir)
            time.sleep(0.1)
            return True
    return False


def get_1st_folder_list(src_dir: str) -> list:
    """获取src_dir下面一级子目录绝对路径列表

    Args:
        src_dir (str): 源文件夹

    Returns:
        list: 一级子目录绝对路径列表
    """
    first_folders_under_src = []
    temp_folder_dir = []
    for dirpath, dirnames, filenames in os.walk(src_dir):
        # print(dirpath)
        temp_folder_dir.append(dirpath)
    try:
        temp_folder_dir.pop(0)  # 把root路径弹出
    except IndexError:
        pass
    # ! 用了很笨的办法获取一级folder的list
    for dir in temp_folder_dir:
        temp = dir.replace(src_dir, '')
        i = 0
        for char in temp:
            if char == "\\":
                i += 1
        if i == 1:
            first_folders_under_src.append(dir)

    return first_folders_under_src


def copy_folder_to_target(src_dir: str, folder_key_name: str, des_dir: str) -> True or False:
    """ 在指定的src目录内遍历一级子目录folder_names，按照关键词找到对应的folder拷贝到des目录

    Args:
        src_dir (str): 源文件夹
        folder_key_names (str): 关键词
        des_dir (str): 目标文件夹，每个生成的变更生成子目录，子目录名要包含原来的说明

    Returns:
        True or False: 成功在src_dir一级子目录folder_names里找到folder_key_name并拷贝到des目录下返回True
        ，否则返回False
    """
    if os.path.exists is False:
        os.mkdir(des_dir)
        print("新建 %s " % des_dir)
    # 遍历生成src_folders下一级子目录绝对路径lists
    first_folders_under_src = get_1st_folder_list(src_dir)

    for dir in first_folders_under_src:
        dir_name_1st = dir.split("\\")[-1]
        if str.find(dir_name_1st, folder_key_name) >= 0:
            target_dir_name = des_dir + "\\" + dir_name_1st
            shutil.copytree(dir, target_dir_name)
            print(dir + " --> " + target_dir_name)
            time.sleep(0.1)
            return True
    return False


def rename_folder_base_on_2nd_folder_name(src_dir: str, key_word: str) -> True or False:
    """根据src_dir下面的文件中的folder_name含有的内容重命名src_dir

    Args:
        src_dir (dir): src_dir下面的一级folder包含说明str，同时需要重命名的文件夹
        key_word (str): 根据key_word找到带说明的folder

    Returns:
        True or False: 如果src_dir下面的一级folder包含说明并重命名src_dir返回True，否则返回False
    """
    folders_name = get_1st_folder_list(src_dir)
    temp_name = []
    for folder in folders_name:
        if str.find(folder, key_word) > 0:
            folder = folder.replace("（", "(")
            folder = folder.replace("）", ")")
            pattern = re.compile('[(](.*?)[)]', re.S)
            temp_name = re.findall(pattern, folder)  # 这个list

    if len(temp_name) == 1:
        print(temp_name[0])
        os.rename(src_dir, src_dir + "--" + key_word + "--" + temp_name[0])


def main():
    # ! main函数中的下面这些代码需要重新整理，目前只能针对PC进行工作，不合理
    # ! 废弃的代码
    # * 处理步骤一：处理xlsx文件，查重生成df
    # 输入数据路径 & 临时文件路径
    file_with_dir = "D:\\CloudStation\\Python\\Project\\CNCC2_DesignChange_Doc\\CNCC2_DesignChanges.xlsx"
    temp_file = "D:\\CloudStation\\Python\\Project\\CNCC2_DesignChange_Doc\\temp.csv"
    # 生成df，存temp文件用于测试，如果df为空退出
    xlsx = base_data_from_xls(file_with_dir)
    df_data = xlsx.get_data_proceed()
    if df_data.empty:
        print("重复编号 in xlsx")
        exit()
    if os.path.exists(temp_file):
        os.remove(temp_file)
    df_data.to_csv(temp_file, encoding="utf_8_sig")

    # * 处理步骤二：处理docx文件，基于模板文件填内容
    # 带索引的base_docx文件 & docx文件输出文件夹
    basic_doc_file = "D:\\CloudStation\\Python\\Project\\CNCC2_DesignChange_Doc\\会展投资合字2019第132号BG_0_--审批表.docx"
    doc_dir = "D:\\CloudStation\\Python\\Project\\CNCC2_DesignChange_Doc\\test\\"

    # 手动填写的索引list用来遍历docx内容进行替换
    # ! 肯定有非常方便的method把df里列名变成list或者一个可迭代obj
    cols_text = ["审批表编号", "提出单位一", "变更原因二", "变更内容三", "何种专业四", "#####"]

    for i in range(0, len(df_data)):
        doc = replace_doc_lines(basic_doc_file, doc_dir)
        doc.make_new_docx(df_data.iloc[i]["审批表编号"])
        for cols_name in cols_text:
            doc.replace_text(cols_name, df_data.iloc[i][cols_name])
        doc.save_proceed_docx()

    # * 处理步骤三：根据excel文件里设计变更编号在指定的folder里找对应的folder拷贝到目标folder里
    # 所有变更存放的绝对路径
    design_change_dirs = ["C:\\Users\\CNCC2-01\\Desktop\\1\\设备\\暖通",
                          "C:\\Users\\CNCC2-01\\Desktop\\1\\设备\\给排水"]
    copy_to_dir = "C:\\Users\\CNCC2-01\\Desktop\\新建文件夹"
    sheet_dir = "D:\\CloudStation\\Python\\Project\\CNCC2_DesignChange_Doc\\test"

    # TODO 建立df_data里面编号和变更的dict，先遍历编号然后遍历变更的两个路径，对应的往对应目录里拷
    df_to_dict = df_data.set_index("变更通知单编号").to_dict()['审批表编号']
    for key, value in df_to_dict.items():
        for dir in design_change_dirs:
            des_dir = copy_to_dir + "\\会展投资合字2019第132号BG" + \
                str(value).rjust(3, '0') + "\\"
            if os.path.exists(des_dir):
                pass
            else:
                os.makedirs(des_dir)
            copy_folder_to_target(dir, key, des_dir)
            copy_file_to_target(sheet_dir, "第132号BG" +
                                str(value).rjust(3, '0'), des_dir)

    # * 处理步骤四：根据生成的审批单folder下的变更文件夹名字中的说明重命名审批单folder
    src_dir = "C:\\Users\\CNCC2-01\\Desktop\\新建文件夹"
    temp_folder_dir = get_1st_folder_list(src_dir)
    df_to_dict = df_data.set_index("变更通知单编号").to_dict()['审批表编号']
    for dir in temp_folder_dir:
        for key, value in df_to_dict.items():
            rename_folder_base_on_2nd_folder_name(dir, key)


def get_design_change_from_xlsx(xlsx_with_dir: str, project_id: str = None) -> pd.DataFrame:
    """处理步骤一：处理xlsx文件，查重生成df

    Args:
        xlsx_with_dir (str): xlsx的路径
        project_name: 具体是哪个地块

    Returns:
        pd.DataFrame: xlsx对应的df
    """
    # * 处理步骤一：处理xlsx文件，查重生成df
    # 输入数据路径 & 临时文件路径
    temp_file = "./temp.csv"
    # 生成df，存temp文件用于测试，如果df为空退出
    xlsx = base_data_from_xls(xlsx_with_dir, project_id)
    df_data = xlsx.get_data_proceed()
    if df_data.empty:
        print("重复编号 in xlsx")
        exit()
    if os.path.exists(temp_file):
        os.remove(temp_file)
    df_data.to_csv(temp_file, encoding="utf_8_sig")
    return df_data


def set_lines_in_doc(df_data: pd.DataFrame, basic_doc_file: str, doc_dir: str, project_id: str):
    """处理步骤二：处理docx文件，基于模板文件填内容

    Args:
        df_data (pd.DataFrame): [从xlsx文件读取的df]
        basic_doc_file (str): [原始doc模板文件的路径]
        doc_dir (str): [批量输出doc文件的路径]
    """

    if project_id == "B25B26":
        # B25B26
        cols_text = xlsx_B25B26_columns
    elif project_id == "B23":
        # B23
        cols_text = xlsx_B23_columns

    for i in range(0, len(df_data)):
        doc = replace_doc_lines(basic_doc_file, doc_dir,
                                df_data.iloc[i]["变更编号和说明"])
        # print(df_data.iloc[i]["审批表编号"])
        doc.make_new_docx(df_data.iloc[i]["审批表编号"])
        for cols_name in cols_text:
            doc.replace_text(cols_name, df_data.iloc[i][cols_name])
        doc.save_proceed_docx()


def CNCC2_DesignChanges(xlsx_with_dir: str, docx_tempalate: str, project_id: str = None, output_dir: str = "./test"):
    df = get_design_change_from_xlsx(xlsx_with_dir, project_id)

    # 带索引的base_docx文件 & docx文件输出文件夹
    set_lines_in_doc(df, docx_tempalate, output_dir, project_id)


def get_file_for_DesignChanges(file_key_word: str, file_type: str):
    """为了根据关键词从上层目录里获取xlsx和docx文件的路径


    Args:
        file_key_word (str): file文件名的关键词
        file_type (str): file扩展名的关键词

    Returns:
        [type]: file文件的完整路径信息
    """
    file_with_dir = ""
    files = os.listdir("./")
    for file in files:
        if file_key_word in file and file_type in file and "$" not in file:
            file_with_dir = os.path.join("./", file)
            print(file_with_dir)

    return file_with_dir


if __name__ == "__main__":

    B25B26_xlsx_with_dir = get_file_for_DesignChanges("B25B26", "xlsx")
    B23_xlsx_with_dir = get_file_for_DesignChanges("B23", "xlsx")

    B25B26_basic_doc_file = get_file_for_DesignChanges("B25B26", "docx")
    B23_basic_doc_file = get_file_for_DesignChanges("B23", "docx")
    CNCC2_DesignChanges(B25B26_xlsx_with_dir,
                        B25B26_basic_doc_file, "B25B26", "./B25B26/")
    # CNCC2_DesignChanges(B23_xlsx_with_dir, B23_basic_doc_file, "B23", "./B23/")
